#!/usr/bin/env python3
"""Render merged-paper.md -> merged-paper.docx in STANDARD academic-journal
formatting (Times New Roman body, black numbered headings, plain tables with
horizontal rules, justified single column, figures embedded with captions
below). Suitable as a base for the EMSE/Springer template.

Figures are referenced in the markdown by italic caption lines containing a
backtick path (e.g. *Figure 1: ... (`stage5_out/fig1...png`)*); we detect the
basename and embed the matching file from figures/, caption below the image."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = Path(__file__).resolve().parent
SRC = _HERE / "merged-paper.md"
DST = _HERE / "merged-paper.docx"
FIGDIR = _HERE.parent / "figures"

BODY = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x55, 0x55, 0x55)

# caption lines now read "*Figure N. ...*"; map the figure NUMBER to its file
FIG_BY_NUM = {
    "1": "fig1_parallel_trends.png",
    "2": "fig2_dynamic_did.png",
    "3": "fig3_transitions.png",
    "4": "fig7_result2_sim.png",
    "5": "fig5_tail_did.png",
    "6": "fig6_size_distortion.png",
    "7": "fig7_axes_forest.png",
}

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = BODY; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.0)
sec.top_margin = sec.bottom_margin = Inches(1.0)

# minimal page-number footer only (standard manuscript)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr"); sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"18")
rpr.append(sz); run.append(rpr); fld.append(run); fp._p.append(fld)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); th = OxmlElement("w:tblHeader"); th.set(qn("w:val"),"true"); trPr.append(th)

GREEK = {r"\rho":"ρ", r"\lambda":"λ", r"\mu":"μ", r"\beta":"β", r"\alpha":"α",
         r"\gamma":"γ", r"\delta":"δ", r"\sigma":"σ", r"\pi":"π", r"\chi":"χ",
         r"\varepsilon":"ε", r"\hat":"", r"\approx":"≈", r"\ge":"≥", r"\le":"≤",
         r"\times":"×", r"\in":"∈", r"\to":"→", r"\sum":"Σ", r"\cdot":"·", r"\,":" "}
SUB = {"0":"₀","1":"₁","2":"₂","3":"₃","4":"₄","5":"₅","6":"₆","7":"₇","8":"₈","9":"₉",
       "+":"₊","-":"₋","i":"ᵢ","j":"ⱼ","k":"ₖ","t":"ₜ"}
SUP = {"2":"²","3":"³"}

def delatex(text):
    """Convert inline LaTeX ($...$) to readable Unicode for a Word manuscript."""
    def conv(m):
        x = m.group(1)
        for a, b in GREEK.items(): x = x.replace(a, b)
        x = re.sub(r"\\text\{([^}]*)\}", r"\1", x)
        x = re.sub(r"\\(?:mathrm|mathbf|mathit)\{([^}]*)\}", r"\1", x)
        # subscripts: _{...} or _x
        x = re.sub(r"_\{([^}]*)\}", lambda mm: "".join(SUB.get(c, c) for c in mm.group(1)), x)
        x = re.sub(r"_([0-9A-Za-z+\-])", lambda mm: SUB.get(mm.group(1), mm.group(1)), x)
        # superscripts: ^{...} or ^x
        x = re.sub(r"\^\{([^}]*)\}", lambda mm: "".join(SUP.get(c, "^"+c) for c in mm.group(1)) if all(c in SUP for c in mm.group(1)) else "^"+mm.group(1), x)
        x = re.sub(r"\^([0-9])", lambda mm: SUP.get(mm.group(1), "^"+mm.group(1)), x)
        x = x.replace("\\", "").replace("{","").replace("}","")
        return x.strip()
    return re.sub(r"\$([^$]+)\$", conv, text)

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
def add_runs(par, text, size=11, italic_all=False):
    text = delatex(text)
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            run = par.add_run(tok[2:-2]); run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = par.add_run(tok[1:-1]); run.font.name = "Courier New"; run.font.size = Pt(size-1)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run = par.add_run(tok[1:-1]); run.italic = True
        else:
            run = par.add_run(tok)
        if not run.font.name: run.font.name = BODY
        if not run.font.size: run.font.size = Pt(size)
        if italic_all: run.italic = True

def _border(el, edge, sz="4", color="000000"):
    e = OxmlElement(f"w:{edge}")
    for k, v in (("w:val","single"),("w:sz",sz),("w:space","0"),("w:color",color)): e.set(qn(k), v)
    el.append(e)

def emit_table(rows):
    """Standard academic table: bold header row, horizontal rules top/under-header/
    bottom only (booktabs-style), no vertical lines, no shading."""
    header = rows[0]
    body = rows[2:] if len(rows) > 1 and set("".join(rows[1]).replace(":","")) <= {"-"," ",""} else rows[1:]
    t = doc.add_table(rows=1, cols=len(header)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = True
    set_repeat_header(t.rows[0])
    for j, c in enumerate(header):
        p = t.rows[0].cells[j].paragraphs[0]; p.text = ""; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, c.strip(), size=10)
        for run in p.runs: run.bold = True
    for row in body:
        cells = t.add_row().cells
        for j, c in enumerate((row + [""]*len(header))[:len(header)]):
            p = cells[j].paragraphs[0]; p.text = ""; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, c.strip(), size=10)
    # booktabs-style: only horizontal rules. Use insideH thin + strong top/bottom.
    tbl = t._tbl; tblPr = tbl.tblPr; borders = OxmlElement("w:tblBorders")
    _border(borders, "top", sz="8")
    _border(borders, "bottom", sz="8")
    _border(borders, "insideH", sz="2", color="888888")
    tblPr.append(borders); doc.add_paragraph()

def maybe_figure(s):
    # match a caption line like "*Figure 3. ...*"
    m = re.match(r"\*Figure\s+(\d+)\.", s)
    if m:
        fname = FIG_BY_NUM.get(m.group(1))
        if fname and (FIGDIR / fname).exists():
            return FIGDIR / fname
    return None

lines = SRC.read_text().splitlines()
i = 0; first_h1 = False; figs = 0; fignum = 0
# strip the author-facing draft note and the references "integrity note" blockquote?
# Keep them — author will clean. We just format faithfully.
while i < len(lines):
    s = lines[i].strip()

    m = re.match(r"!\[.*?\]\((.+?)\)", s)
    if m and Path(m.group(1)).exists():
        doc.add_picture(m.group(1), width=Inches(5.5)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1; continue

    figpath = maybe_figure(s)
    if figpath is not None:
        doc.add_picture(str(figpath), width=Inches(5.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(cap, s, size=9.5, italic_all=True)
        for run in cap.runs: run.font.color.rgb = GREY
        figs += 1; i += 1; continue

    if s.startswith("|") and "|" in s[1:]:
        block = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            block.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
        emit_table(block); continue

    if s.startswith("#"):
        lvl = len(s) - len(s.lstrip("#")); txt = s.lstrip("#").strip()
        if lvl == 1 and not first_h1:
            first_h1 = True
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt); run.bold = True; run.font.size = Pt(16); run.font.name = BODY; run.font.color.rgb = BLACK
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(12)
        else:
            p = doc.add_paragraph(); size = {2:13,3:11.5,4:11}.get(lvl,11)
            run = p.add_run(txt); run.bold = True; run.font.size = Pt(size); run.font.name = BODY; run.font.color.rgb = BLACK
            p.paragraph_format.space_before = Pt(12 if lvl==2 else 8); p.paragraph_format.space_after = Pt(4)
        i += 1; continue

    if s == "---": i += 1; continue

    if s.startswith(">"):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.right_indent = Inches(0.2)
        add_runs(p, s.lstrip("> ").strip(), size=10, italic_all=True)
        i += 1; continue

    if re.match(r"^\d+\.\s", s):
        p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s","",s).strip())
        i += 1; continue
    if s.startswith("- ") or s.startswith("* "):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, s[2:].strip())
        i += 1; continue

    if not s: i += 1; continue

    # display equation: a line that is entirely $...$  -> convert, center, italic
    if s.startswith("$$") or (s.startswith("$") and s.endswith("$") and s.count("$") == 2):
        eq = delatex(s.strip("$").strip())   # delatex needs the $ wrappers, so wrap then strip
        if "$" not in eq:  # if delatex didn't fire (no inner $), convert manually
            eq = delatex("$" + s.strip("$").strip() + "$")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(eq); run.italic = True; run.font.name = BODY; run.font.size = Pt(11)
        i += 1; continue

    p = doc.add_paragraph()
    if s.startswith("**Somil Sharma**") or s.startswith("Independent Researcher") or s.startswith("*Draft for submission"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, s, size=11)
    elif s.startswith("**Table"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, s, size=10)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; add_runs(p, s, size=11)
    i += 1

doc.save(DST)
print(f"wrote {DST}  ({figs} figures embedded)")
